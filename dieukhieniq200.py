import serial
import time
import pdb
import pyvisa
import os
import subprocess
import glob
import numpy as np
import re
from docx import Document
from datetime import datetime

  
rm = pyvisa.ResourceManager('@py')
# Kết nối hai máy đo


ip_1 = "TCPIP0::169.254.207.15::inst0::INSTR"
ip_2 = "TCPIP0::169.254.207.25::inst0::INSTR"

# Kết nối máy phát và máy thu
try:
    sig_anl = rm.open_resource(ip_1)
    print("✅ Kết nối thành công!")
    print("Thông tin máy đo:", sig_anl.query("*IDN?"))
except pyvisa.VisaIOError as e:
    print("❌ Lỗi kết nối:", e)
try:
    sig_gen = rm.open_resource(ip_2)
    print("✅ Kết nối thành công!")
    print("Thông tin máy đo:", sig_gen.query("*IDN?"))
except pyvisa.VisaIOError as e:
    print("❌ Lỗi kết nối:", e)

# Lệnh để giết tất cả tiến trình sử dụng cổng /dev/tty.usbserial-120
def kill_process():
    command = "kill -9 $(lsof -t /dev/tty.usbserial-120)"

    try:
        # Thực thi lệnh shell
        subprocess.run(command, shell=True, check=True)
        print("✅ Đã tắt các tiến trình đang sử dụng cổng /dev/tty.usbserial-120.")
    except subprocess.CalledProcessError as e:
        print(f"❌ Lỗi khi thực thi lệnh: {e}")
    except Exception as e:
        print(f"❌ Lỗi không xác định: {e}")
kill_process()


# Cấu hình cổng Serial


def choose_serial_port():
    # Liệt kê các cổng serial có dạng /dev/tty.*
    ports = glob.glob('/dev/tty.*')

    if not ports:
        print("⚠️ Không tìm thấy cổng serial nào.")
        return None

    # Hiển thị danh sách cổng
    print("🔌 Danh sách cổng serial có sẵn:")
    for idx, port in enumerate(ports, 1):
        print(f"{idx}. {port}")

    # Chọn cổng
    while True:
        try:
            choice = int(input("👉 Nhập số thứ tự cổng bạn muốn dùng: "))
            if 1 <= choice <= len(ports):
                return ports[choice - 1]
            else:
                print("❌ Lựa chọn không hợp lệ, vui lòng thử lại.")
        except ValueError:
            print("❌ Vui lòng nhập số!")

# Chạy hàm chọn cổng
port_name = choose_serial_port()

if port_name:
    baud_rate = 115200  # Tốc độ baud IQ200
    iq200 = serial.Serial(port_name, baud_rate, timeout=5, rtscts=False, dsrdtr=False)
    print(f"✅ Đã kết nối đến {port_name} với baudrate {baud_rate}")
else:
    print("⛔ Không có cổng nào được chọn. Thoát chương trình.")


def send_and_wait(command, expect_prompt, timeout=2, retries=20):
    for attempt in range(retries):
        iq200.reset_input_buffer()
        iq200.reset_output_buffer()
        iq200.write(command)
        iq200.flush()
        start_time = time.time()
        response = ''
        response_line = ''
        while time.time() - start_time < timeout:
            if iq200.in_waiting:
                response_line = iq200.readline().decode('ascii', errors='ignore')
                # print(response_line)
                response += response_line
                if expect_prompt in response_line:
                    print(f"Sent: {command.strip()}, Received: {response.strip()}")
                    return response
            time.sleep(0.1)
        print(f"Đang kiểm tra {command} lần thứ {attempt + 1} trả về {response}. Retrying...")
    raise ValueError(f"Expected prompt '{expect_prompt}' not found in response after {retries} attempts.")


    #Đăng nhập IQ200
# Đăng nhập vào IQ200
def iq200_login():
    skip = input("👉 Bỏ qua đăng nhập modem? (Y/N): ").strip().lower()
    if skip == 'y':
        print("⏩ Bỏ qua bước đăng nhập.")
        return True

    try:
        iq200.reset_input_buffer()
        iq200.reset_output_buffer()
        send_and_wait(b"\r\n", expect_prompt="iq-linux login:")
        time.sleep(1)
        send_and_wait(b"root\n", expect_prompt="Password:")
        time.sleep(1)
        iq200.write(b"P@55w0rd!\n")
        time.sleep(1)
        send_and_wait(b"mewsh\n", expect_prompt="[Remote]$")
        time.sleep(1)
        send_and_wait(b"\r\n", expect_prompt="[Remote]$")
        print("✅ Đăng nhập thành công vào IQ200")
        return True

    except Exception as e:
        print(f"❌ Đăng nhập thất bại: {e}")
        retry = input("🔄 Bạn có muốn thử đăng nhập lại không? (y/n): ").strip().lower()
        if retry == 'y':
            return iq200_login()
        else:
            print("⏹️ Dừng đăng nhập.")
            return False

results_tx = []
results_rx = []
results_10M = []
def iq200_serial():
    global sn_iq200

    try:
        iq200.reset_input_buffer()
        iq200.reset_output_buffer()

        # Gửi lệnh trắng để "thức" modem
        send_and_wait(b"\r\n", expect_prompt="[Remote]$")
        
        # Gửi lệnh DID và chờ prompt [Remote]$
        send_and_wait(b"DID\r\n", expect_prompt="[Remote]$")
        time.sleep(0.5)

        # Đọc phản hồi
        response = iq200.read_all().decode('ascii', errors='ignore')

        for line in response.splitlines():
            if "SN:" in line:
                # Loại bỏ phần sau dấu "#" nếu có
                line = line.split("#")[0]

                serial_match = re.search(r"SN:\s*(\d+)", line)
                if serial_match:
                    sn_iq200 = serial_match.group(1)
                    print(f"🔢 Serial Number: {sn_iq200}")
                    return sn_iq200

        print("❌ Không tìm thấy Serial Number trong phản hồi.")
        return None

    except Exception as e:
        print(f"🚨 Lỗi khi lấy Serial Number: {e}")
        return None

# Kiểm tra tần số thu IQ200
def test_rx(iq200):
    global results_rx  
    sig_gen.write("SYST:PRES")
    time.sleep(1)

    try:
        while True:
            freq_input = input("Nhập TẦN SỐ THU (MHz) [start, stop, step] (Enter = mặc định 950,2150,600): ")
            if not freq_input.strip():  # Dùng mặc định nếu bỏ trống
                freq_list = np.arange(950, 2150 + 0.01, 600)
                break
            try:
                parts = [x.strip() for x in freq_input.split(",")]
                if len(parts) != 3:
                    raise ValueError("⚠️ Bạn cần nhập đúng 3 số: start,stop,step")
                start_freq, stop_freq, step_freq = map(float, parts)
                if step_freq <= 0:
                    raise ValueError("⚠️ Step phải là số dương.")
                freq_list = np.arange(start_freq, stop_freq + step_freq/2, step_freq)
                break
            except ValueError as ve:
                print(f"❌ {ve}")

            print(f"Tần số kiểm tra: {list(freq_list)}")


        for frequency_rx in freq_list:
            print(f"\n📶 Đang kiểm tra tại tần số thu: {frequency_rx} MHz")

            # Cấu hình máy phát
            sig_gen.write(f"FREQ {frequency_rx} MHz")
            sig_gen.write("OUTP ON")
            sig_gen.write("POW -10 dBm")
            sig_gen.write("OUTP:MOD:STAT OFF")
            time.sleep(1)  # Đợi máy phát tín hiệu ổn định

            response1 = ""
            while response1.find("- Set") == -1:
                send_and_wait(b"\r\n", expect_prompt="[Remote]$")
                set_freq_iq200 = f"rx_freq {frequency_rx * 1000000}\r\n"
                iq200.reset_input_buffer()
                iq200.write(set_freq_iq200.encode('ascii'))
                iq200.flush()
                time.sleep(2)
                response1 = iq200.read_all().decode('ascii', errors='ignore')
                time.sleep(2)

            response = ""
            while response.find("Waiting for Demod Lock") == -1:
                iq200.write(b"rx\r\n")
                iq200.flush()
                time.sleep(2)   
                response = iq200.read_all().decode('ascii', errors='ignore')
                time.sleep(2)
                iq200.reset_input_buffer()

            # Phân tích dữ liệu phản hồi
            response_lines = response.splitlines()
            frequency_rx_results = None
            power_rx_results = None
            for line in response_lines:
                if "frequency:" in line:
                    frequency_rx_results = float(line.split(":")[1].strip())
                if "power:" in line:
                    power_rx_results = float(line.split(":")[1].split("#")[0].strip())
                    break

            if frequency_rx_results is not None and power_rx_results is not None:
                results_rx.append([float(frequency_rx), float(frequency_rx_results), float(power_rx_results)])
            else:
                print(f"❌ Không thể lấy dữ liệu tần số hoặc công suất cho tần số {frequency_rx} MHz")

        sig_gen.write("OUTP OFF")
        print("📊 Kết quả kiểm tra RX:")
        for freq_gen,freq_rx, power in results_rx:
            print(f"- Tần số kiểm tra:{freq_gen} MHz  | Tần số thu được {freq_rx} MHz | Công suất: {power} dBm")

    except Exception as e:
        print(f"🚨 Lỗi trong quá trình kiểm tra RX: {e}")


# Kiểm tra tần số phát IQ200
def test_tx(iq200):
    global results_tx 

    sig_anl.write("SYST:PRES")
    sig_anl.write("INST:SEL SA")
    sig_anl.write("SYST:PRES")
    time.sleep(1)

    try:
        # Nhập và kiểm tra dải tần số phát
        while True:
            freq_input = input("Nhập TẦN SỐ PHÁT (MHz) [start, stop, step] (Enter = mặc định 950-2400-725): ")
            if not freq_input.strip():
                freq_list = np.arange(950, 2400 + 0.01, 725)  # dùng numpy cho phép số thực
                break
            try:
                start_freq, stop_freq, step_freq = map(float, freq_input.split(","))
                freq_list = np.arange(start_freq, stop_freq + step_freq/2, step_freq)  # tránh lệch tròn
                break
            except ValueError:
                print("⚠️ Sai định dạng! Nhập theo: start,stop,step (VD: 950,2400,362.5) hoặc Enter để dùng mặc định.")

        # Nhập và kiểm tra dải công suất phát
        while True:
            power_input = input("Nhập CÔNG SUẤT PHÁT (dBm) [start, stop, step] (Enter = mặc định 0,-35,-17.5): ")
            if not power_input.strip():
                power_list = np.arange(0, -35 + 0.01, -17.5)
                break
            try:
                start_power, stop_power, step_power = map(float, power_input.split(","))
                power_list = np.arange(start_power, stop_power + step_power/2, step_power)
                break
            except ValueError:
                print("⚠️ Sai định dạng! Nhập theo: start,stop,step (VD: 0,-35,-7) hoặc Enter để dùng mặc định.")


        sig_anl.write("SYST:PRES")
        time.sleep(1)

        for frequency_tx in freq_list:
            print(f"\n📶 Đang kiểm tra tại tần số phát: {frequency_tx} MHz")

            for power_level in power_list:
                print(f"⚡ Đặt công suất: {power_level} dBm")

                sig_anl.write("SYST:PRES")
                send_and_wait(b"\r\n", expect_prompt="[Remote]$")

                set_freq_iq200 = f"tx_cross_pol_test cw {frequency_tx} {power_level * 10} 1 16 1 3\r\n"
                iq200.reset_input_buffer()
                iq200.reset_output_buffer()
                iq200.write(set_freq_iq200.encode('ascii'))
                iq200.flush()
                time.sleep(2)

                early_response = iq200.read_all().decode('ascii', errors='ignore')

                if "Cross polarization test on" not in early_response:
                    time.sleep(8)
                    response1 = iq200.read_all().decode('ascii', errors='ignore')
                else:
                    response1 = early_response

                if "Cross polarization test on" in response1:
                    print("✅ Modem xác nhận đã bật Cross Polarization Test.")
                else:
                    print("❌ Modem phản hồi sai, lệnh có thể không chạy đúng.")

                sig_anl.write("FREQ:START 100 MHz")
                sig_anl.write("FREQ:STOP 11 GHz")
                time.sleep(1)
                sig_anl.write("CALC:MARK:MAX")
                time.sleep(1)

                peak_center = sig_anl.query("CALC:MARK:X?")
                sig_anl.write(f"SENS:FREQ:CENT {peak_center}")
                sig_anl.write("SENS:FREQ:SPAN 100 MHz")
                time.sleep(1)
                sig_anl.write("CALC:MARK:MAX")
                time.sleep(1)
                sig_anl.write("CALC:MARK:CENT")
                time.sleep(1)
                sig_anl.write("SENS:FREQ:SPAN 10 MHz")
                time.sleep(1)
                sig_anl.write("CALC:MARK:MAX")
                time.sleep(1)
                sig_anl.write("CALC:MARK:CENT")
                time.sleep(1)

                peak_freq = round(float(sig_anl.query("CALC:MARK:X?")) / 1_000_000, 4)
                peak_amp = round(float(sig_anl.query("CALC:MARK:Y?")), 4)

                print(f"🎯 Peak Freq: {peak_freq} MHz | 🎯 Peak AMP: {peak_amp} dBm")

                results_tx.append((float(frequency_tx), float(power_level), float(peak_freq), float(peak_amp)))

        print("\n📊 Kết quả kiểm tra TX:")
        for freq, power, peak_f, peak_a in results_tx:
            print(f"- Tần số test: {freq} MHz | Công suất: {power} dBm | Peak: {peak_f} MHz | Amplitude: {peak_a} dBm")

    except Exception as e:
        print(f"🚨 Lỗi trong quá trình kiểm tra: {e}")


# Kiểm tra 10 MHz IQ200
def test_10M(iq200):
    global results_10M
    sig_anl.write("SYST:PRES")
    time.sleep(1)

    try:
        sig_anl.write("INST:SEL PNOISE")
        time.sleep(1)

        sig_anl.write("INIT:LPL")
        time.sleep(1)

        sig_anl.write("CALC:LPL:DEC:TABL ON")  # Bật Decade Table
        time.sleep(1)

        sig_anl.write("SENS:FREQ:CARR:SEAR")  # Tìm carrier
        time.sleep(1)
        sig_anl.query("*OPC?")
        time.sleep(1)
        sig_anl.write(":SENSE:LPL:FREQ:OFFS:STAR 100")
        sig_anl.write(":SENSE:LPL:FREQ:OFFS:STOP 10E6")
        sig_anl.query("*OPC?")
        time.sleep(20)
        # === Thêm phần đặt marker tại các offset ===
        offsets = [1000, 10000, 100000]  # Offset marker
        for idx, offset in enumerate(offsets, start=1):
            sig_anl.write(f":CALC:LPL:MARK{idx}:MODE POS")
            sig_anl.write(f":CALC:LPLot:MARK{idx}:X {offset}")
            sig_anl.write(f":CALC:LPL:MARK{idx}:TRAC 1")
            time.sleep(1)

            # Lấy giá trị và lưu vào list
            marker_x = float(sig_anl.query(f":CALC:LPL:MARK{idx}:X?"))
            marker_y = float(sig_anl.query(f":CALC:LPL:MARK{idx}:Y?"))
            results_10M.append((marker_x, marker_y))

        # In kết quả ra màn hình
        print("Offset (Hz) | Phase Noise (dBc/Hz)")
        print("-" * 30)
        for x, y in results_10M:
            print(f"{int(x):>10} | {y:>8.2f}")

    except Exception as e:
        print(f"🚨 Lỗi trong quá trình kiểm tra: {e}")
# Sử lý kết quả
def save_results_to_word():
    global results_tx, results_rx, results_10M, sn_iq200

    # Tạo document mới
    doc = Document()
    doc.add_heading('Kết quả kiểm tra IQ200', level=1)
    doc.add_paragraph(f"Ngày kiểm tra: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Kết quả TX
    doc.add_heading('Kết quả kiểm tra TX', level=2)
    for freq, power, peak_freq, peak_amp in results_tx:
        doc.add_paragraph(f"Tần số test: {freq} MHz | Công suất: {power} dBm | Peak: {peak_freq} MHz | Amplitude: {peak_amp} dBm")

    # Kết quả RX
    doc.add_heading('Kết quả kiểm tra RX', level=2)
    for freq, freq_rx, power in results_rx:
        doc.add_paragraph(f"Tần số: {freq} MHz | Công suất: {power} dBm")

    # Kết quả 10 MHz
    doc.add_heading('Kết quả kiểm tra 10 MHz', level=2)
    for x, y in results_10M:
        doc.add_paragraph(f"Offset (Hz): {int(x)} | Phase Noise (dBc/Hz): {y:.2f}")

    # Đường dẫn lưu file
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    output_folder = os.path.join(desktop_path, "kiemtraiq200", "ketquakiemtra")
    os.makedirs(output_folder, exist_ok=True)

    file_name = f"kqktiq200_{sn_iq200}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = os.path.join(output_folder, file_name)

    # Lưu file
    doc.save(file_path)
    print(f"✅ Kết quả đã được lưu vào file: {file_path}")

def replace_placeholder_in_paragraphs(doc, placeholder, replacement_text):
    """Thay thế placeholder trong các đoạn văn (paragraphs)."""
    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = para.text.replace(placeholder, replacement_text)

def replace_placeholder_in_tables(doc, placeholder, replacement_text):
    """Thay thế placeholder trong các ô bảng."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, replacement_text)

def save_results_to_template(results_tx, results_rx, results_10M):
    # Tạo thư mục
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    base_folder = os.path.join(desktop_path, "kiemtraiq200")
    template_folder = os.path.join(base_folder, "template")
    output_folder = os.path.join(base_folder, "ketquakiemtra")

    os.makedirs(template_folder, exist_ok=True)
    os.makedirs(output_folder, exist_ok=True)

    template_path = os.path.join(template_folder, "templatekiemtraiq200.docx")
    if not os.path.isfile(template_path):
        print(f"⚠️ Không tìm thấy template: {template_path}")
        return

    doc = Document(template_path)

    # Tạo nội dung thay thế
    now = datetime.now()
    timestamp = datetime.now().strftime("%m%d%Y_%H%M")
    day = str(now.day)
    month = str(now.month)
    year = str(now.year)

    # Kiểm tra TX
    tx_text_freq = "Đạt"
    tx_text_amp = "Đạt"
    for freq, power, peak_freq, peak_amp in results_tx:
        if abs(freq - peak_freq) >= 0.0001:
            tx_text_freq = "Không đạt"
        if abs(power - peak_amp) >= 5:
            tx_text_amp = "Không đạt"

    # Kiểm tra RX
    rx_text_freq = "Đạt"
    for freq, freq_rx, amp in results_rx:
        if abs(freq - freq_rx) >= 0.0001 or abs(amp + 10) >= 5:
            rx_text_freq = "Không đạt"
            break

    # Ghi phase noise vào bảng
    if results_10M and doc.tables:
        table = doc.tables[1]  # Giả sử bảng đầu tiên là nơi ghi phase noise
        for index in range(min(len(results_10M), len(table.rows) * len(table.columns))):
            amp = results_10M[index][1]  # Chỉ lấy biên độ (dBc/Hz)
            row = (index % 4) + 5       # Dòng bắt đầu từ 1 (bỏ qua tiêu đề)
            col = 5
            if row < len(table.rows) and col < len(table.columns):
                table.cell(row, col).text = f"{amp:.2f}"
    # Thay thế placeholders
    for placeholder, replacement in {
        "{{day}}": day,
        "{{month}}": month,
        "{{year}}": year,
        "{{tx_text_freq}}": tx_text_freq,
        "{{tx_text_amp}}": tx_text_amp,
        "{{rx_text_freq}}": rx_text_freq,
        "{{sn_iq200}}": sn_iq200,
    }.items():
        replace_placeholder_in_paragraphs(doc, placeholder, replacement)
        replace_placeholder_in_tables(doc, placeholder, replacement)

    # Lưu kết quả
    output_filename = f"pktiq200_{sn_iq200}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = os.path.join(output_folder, output_filename)
    doc.save(output_path)

    print(f"✅ Đã lưu kết quả vào: {output_path}")


def display_results():
    global results_tx, results_rx, results_10M
    print("📊 Kết quả kiểm tra đã được lưu vào file kết quả.")
    print(results_tx)
    print(results_rx)
    print(results_10M)
    print("✅ Hoàn tất quá trình kiểm tra.")

def main():
    global results_tx, results_rx, results_10M
    while True:
        if iq200_login():
            break
        else:
            print("🔑 Đăng nhập không thành công. Vui lòng thử lại...\n")
    iq200_serial()
    # Vòng lặp kiểm tra
    while True:
        # Thực hiện các bài test]
        results_tx= []
        results_rx = []
        results_10M = []
        test_tx(iq200)
        test_rx(iq200)
        test_10M(iq200)
        display_results()
        save_results_to_word()
        save_results_to_template(results_tx, results_rx, results_10M)

        # Hỏi người dùng có muốn kiểm tra lại không
        retry = input("\n🔁 Bạn có muốn kiểm tra lại không? (y/n): ").strip().lower()
        if retry != 'y':
            print("\n🎉 Quá trình kiểm tra đã kết thúc. Cảm ơn bạn đã sử dụng!")
            break
main()



